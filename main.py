# ATS_resume_cloner.py
import gradio as gr
from docx import Document
from docx.shared import Pt, RGBColor
import fitz
import io
import re
import os
import tempfile
from groq import Groq
from dotenv import load_dotenv

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))        # ← put your key in environment
# Model `llama-3.1-70b-versatile` has been decommissioned.
# Allow overriding via env var and fall back to a currently supported model.
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")


def read_any_resume(file):
    """
    Read a resume coming from Gradio's `gr.File` component.

    Newer Gradio versions often pass a `NamedString` (a path-like string)
    instead of a file object with `.read()`, which caused the
    `'NamedString' object has no attribute 'read'` error.

    This helper normalizes the input to a file path and then branches
    on the extension.
    """
    if file is None:
        return ""

    # Normalize to a filesystem path
    if isinstance(file, bytes):
        path = file.decode("utf-8")
    elif isinstance(file, str):
        path = file
    elif hasattr(file, "name") and isinstance(file.name, str):
        # For Gradio's NamedString or file-like objects with a .name attribute
        path = file.name
    else:
        # Fallback: treat it as a binary stream
        content = file.read()
        return content.decode("utf-8", errors="ignore")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        # Let PyMuPDF open directly from the path
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    elif ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore")

def extract_with_llama70b(text):
    prompt = f"""
    Extract the resume in this exact JSON format. Return ONLY valid JSON.

    {{
      "name": "",
      "location": "",
      "email": "",
      "phone": "",
      "linkedin": "",
      "github": "",
      "summary": "",
      "skills": "",
      "experience": [
        {{
          "title": "",
          "company": "",
          "dates": "",
          "location": "",
          "bullets": [""]
        }}
      ],
      "education": [
        {{"degree": "", "school": "", "year": ""}}
      ]
    }}

    Resume text:
    {text[:16000]}
    """

    chat = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=3000
    )
    m = re.search(r"\{.*\}", chat.choices[0].message.content, re.DOTALL)
    import json
    return json.loads(m.group())

def generate_summary_from_resume(text, experience_data, education_data, skills_data):
    """
    Generate a structured professional summary in the format of main_resume.docx.
    Creates categories like: Primary Roles, Backend, Frontend, Database, Cloud/DevOps, etc.
    """
    # Build context from extracted data
    exp_text = ""
    if experience_data:
        for exp in experience_data[:3]:  # Use top 3 experiences
            exp_text += f"{exp.get('title', '')} at {exp.get('company', '')} ({exp.get('dates', '')}). "
    
    skills_list = []
    if isinstance(skills_data, list):
        skills_list = [str(s).strip() for s in skills_data if s]
    elif skills_data:
        skills_list = [s.strip() for s in str(skills_data).split(",") if s.strip()]
    
    skills_text = ", ".join(skills_list[:20])  # Top 20 skills for context
    
    prompt = f"""
    Based on the following candidate information, generate a structured resume summary in this EXACT format (each line on a new line):
    
    Primary Roles: [list 2-3 primary roles like "Data Scientist, Computer Vision Engineer, ML Engineer"]
    Backend: [list backend technologies, e.g., "Python, Java, Spark, Hadoop"]
    Frontend: [list frontend technologies if any, or "N/A"]
    API: [list API types, e.g., "REST, GraphQL"]
    Database: [list databases, e.g., "MySQL, MongoDB, Cassandra, Oracle"]
    Cloud/DevOps: [list cloud platforms and DevOps tools, e.g., "AWS, Azure, Docker, Kubernetes"]
    DevOps: [list additional DevOps tools, e.g., "Jenkins, Terraform, GitLab"]
    Leadership: [mention leadership/mentoring if applicable, or "N/A"]
    Industry: [mention industries if available, or "N/A"]
    
    Experience context: {exp_text}
    Skills available: {skills_text}
    
    IMPORTANT:
    - Categorize skills appropriately
    - Use the exact format above with colons
    - Each category on a separate line
    - Only include categories that have relevant skills
    - Be concise (3-5 items per category max)
    - Return ONLY the formatted text, no additional explanation
    """
    
    try:
        chat = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=400
        )
        summary = chat.choices[0].message.content.strip()
        # Clean up any quotes or extra formatting
        summary = summary.strip('"').strip("'").strip()
        return summary
    except Exception as e:
        # If generation fails, return empty string
        return ""

def apply_ATS_template(template_bytes, data):
    doc = Document(io.BytesIO(template_bytes))

    # === 1. Header (Name + contact line) ===
    doc.paragraphs[0].runs[0].text = data["name"]
    doc.paragraphs[1].runs[0].text = f"{data['location']} | Email: {data['email']} | Phone {data['phone']}"

    # === 2. Summary ===
    summary_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("SUMMARY"):
            summary_idx = i
            break
    if summary_idx is not None:
        summary_para = doc.paragraphs[summary_idx]
        
        # Find the end of the summary block (before the next major heading)
        end_idx = summary_idx + 1
        while end_idx < len(doc.paragraphs):
            text = doc.paragraphs[end_idx].text.strip()
            if not text:
                end_idx += 1
                continue
            # Stop at next major section
            if any(h in text for h in ("PORTFOLIO", "WORK AUTHORIZATION", "SKILL MATRIX", "EDUCATION", "WORK EXPERIENCE")):
                break
            end_idx += 1
        
        # Remove all old summary content paragraphs (everything after the heading)
        for p in doc.paragraphs[summary_idx + 1:end_idx]:
            p._element.getparent().remove(p._element)
        
        # Clear and rewrite the SUMMARY heading paragraph
        summary_para.clear()
        r = summary_para.add_run("SUMMARY")
        r.bold = True
        r.font.size = Pt(10)
        
        # Add the new summary text - handle structured format with multiple lines
        if data.get("summary"):
            summary_text = str(data["summary"]).strip()
            if summary_text:
                # Split by newlines to handle structured format (each category on separate line)
                summary_lines = [line.strip() for line in summary_text.split('\n') if line.strip()]
                insert_pos = summary_idx + 1
                
                for line in summary_lines:
                    # Insert each line as a separate paragraph
                    if insert_pos < len(doc.paragraphs):
                        anchor = doc.paragraphs[insert_pos]
                        new_p = anchor.insert_paragraph_before(line)
                    else:
                        new_p = doc.add_paragraph(line)
                    # Set font size for summary text
                    for run in new_p.runs:
                        run.font.size = Pt(10)
                    insert_pos += 1

    # === 3. Portfolio links ===
    portfolio_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("PORTFOLIO"):
            portfolio_idx = i
            break
    if portfolio_idx is not None:
        portfolio_para = doc.paragraphs[portfolio_idx]

        # Find the end of the portfolio block (before the next major heading)
        end_idx = portfolio_idx + 1
        while end_idx < len(doc.paragraphs):
            text = doc.paragraphs[end_idx].text.strip()
            if not text:
                end_idx += 1
                continue
            if any(h in text for h in ("WORK AUTHORIZATION", "SKILL MATRIX", "EDUCATION", "WORK EXPERIENCE")):
                break
            end_idx += 1

        # Remove old portfolio lines (everything after the heading up to end_idx)
        for p in doc.paragraphs[portfolio_idx + 1:end_idx]:
            p._element.getparent().remove(p._element)

        # Rewrite the heading paragraph itself
        portfolio_para.clear()
        r = portfolio_para.add_run("PORTFOLIO")
        r.font.size = Pt(10)
        r.bold = True

        # Add new link lines as separate, non-bulleted paragraphs
        insert_pos = portfolio_idx + 1
        # Helper to insert a paragraph after portfolio heading, preserving order
        def _insert_after_portfolio(text):
            nonlocal insert_pos
            if insert_pos < len(doc.paragraphs):
                anchor = doc.paragraphs[insert_pos]
                new_p = anchor.insert_paragraph_before(text)
            else:
                new_p = doc.add_paragraph(text)
            insert_pos += 1
            return new_p

        if data.get("linkedin"):
            _insert_after_portfolio(f"LinkedIn: {data['linkedin']}")
        if data.get("github"):
            _insert_after_portfolio(f"GitHub: {data['github']}")

    # === 4. Skill Matrix ===
    # Use LLM to analyze resume and create appropriate skill category headers
    # Then generate descriptive bullet points for each category
    skills_raw = data.get("skills", "")
    if isinstance(skills_raw, list):
        skills_list = [str(s).strip() for s in skills_raw if s]
    else:
        skills_text = str(skills_raw).strip() if skills_raw else ""
        skills_list = [s.strip() for s in skills_text.split(",") if s.strip()] if skills_text else []
    
    if skills_list:
        # First, analyze skills and experience to determine appropriate category headers
        experience_context = ""
        if data.get("experience"):
            for exp in data.get("experience", [])[:3]:
                title = exp.get('title', '')
                company = exp.get('company', '')
                experience_context += f"{title} at {company}. "
        
        skills_text_for_prompt = ", ".join(skills_list)
        
        # Step 1: Generate appropriate category headers based on skills and experience
        header_prompt = f"""
        Analyze the following candidate's skills and experience, then suggest 3-5 appropriate category headers for organizing their skills in a resume skill matrix.
        
        Skills: {skills_text_for_prompt}
        Experience: {experience_context}
        
        Based on the skills and experience, create category headers that logically group the skills. 
        Examples of good headers:
        - "Application/Software Development" (for programming languages, frameworks)
        - "Database/SQL/Relational Database/NoSQL" (for databases)
        - "Cloud/AWS/DevOps" (for cloud platforms, DevOps tools)
        - "Machine Learning/Data Science" (for ML/AI skills)
        - "Frontend Technologies" (for frontend frameworks)
        - "Tools/IDE/Editors" (for development tools)
        
        Return ONLY a comma-separated list of 3-5 category header names. Each header should be descriptive and professional.
        Example format: "Application/Software Development, Database/SQL/Relational Database/NoSQL, Cloud/AWS/DevOps, Tools/IDE/Editors"
        """
        
        try:
            header_chat = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": header_prompt}],
                temperature=0.3,
                max_tokens=200
            )
            headers_response = header_chat.choices[0].message.content.strip()
            # Parse headers (remove quotes, split by comma)
            headers_response = headers_response.strip('"').strip("'").strip()
            category_headers = [h.strip() for h in headers_response.split(",") if h.strip()]
            
            # If LLM didn't provide good headers, use defaults
            if not category_headers or len(category_headers) < 2:
                category_headers = ["Application/Software Development", "Database/SQL/Relational Database/NoSQL", 
                                   "Cloud/AWS/DevOps", "Tools/IDE/Editors"]
        except Exception:
            # Fallback to default headers
            category_headers = ["Application/Software Development", "Database/SQL/Relational Database/NoSQL", 
                               "Cloud/AWS/DevOps", "Tools/IDE/Editors"]
        
        # Step 2: Generate skill matrix content with the determined headers
        headers_text = ", ".join(category_headers)
        prompt = f"""
        Based on the candidate's skills and experience, generate a professional skill matrix with the following category headers:
        {headers_text}
        
        For each category, create 3-5 descriptive bullet points in professional resume style.
        Format:
        
        {category_headers[0]}:
        - [bullet point describing experience with relevant skills]
        - [bullet point 2]
        - [bullet point 3]
        
        {category_headers[1] if len(category_headers) > 1 else "Category 2"}:
        - [bullet point 1]
        - [bullet point 2]
        - [bullet point 3]
        
        [Continue for all categories]
        
        Skills available: {skills_text_for_prompt}
        Experience context: {experience_context}
        
        CRITICAL REQUIREMENTS:
        - Each category header must appear EXACTLY ONCE
        - Use the EXACT category header names provided above
        - Write in professional resume style: use action verbs, be direct and concise
        - NEVER use phrases like "The candidate has...", "The candidate is...", "As a...", "Having worked as...", "The candidate demonstrated..."
        - Use direct statements like: "Experience in...", "Proficient in...", "Developed...", "Implemented...", "Skilled in...", "Expertise in...", "Knowledge of...", "Familiarity with..."
        - Example GOOD bullets:
          * "Experience in building scalable web applications using Python, Django, and JavaScript"
          * "Proficient in database design and optimization with MySQL, PostgreSQL, and MongoDB"
          * "Skilled in deploying cloud infrastructure using AWS EC2, S3, and Lambda"
        - Example BAD bullets (DO NOT USE):
          * "The candidate has experience in Python" ❌
          * "As a developer, the candidate worked with Django" ❌
          * "Having worked at Company X, the candidate developed..." ❌
        - Each bullet should be a complete sentence describing experience/expertise
        - Naturally incorporate skill names into descriptive sentences
        - Use "- " prefix for each bullet point
        - Do NOT repeat category headers
        - Write as if describing your own experience (implied first person, no pronouns)
        - Return ONLY the formatted text with category headers (once each) and bullets
        """
        
        try:
            chat = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=800
            )
            skill_matrix_content = chat.choices[0].message.content.strip()
            
            # Parse and insert the skill matrix content
            # Find the first category header or "SKILL MATRIX" to locate where to start
            skill_matrix_start_idx = None
            for i, para in enumerate(doc.paragraphs):
                # Check if this paragraph matches the first generated category header
                para_text = para.text.strip()
                if category_headers and category_headers[0] in para_text:
                    skill_matrix_start_idx = i
                    break
                # Fallback: look for "SKILL MATRIX" or "Application/Software Development"
                if "SKILL MATRIX" in para_text.upper() or "Application/Software Development" in para_text:
                    skill_matrix_start_idx = i + 1 if "SKILL MATRIX" in para_text.upper() else i
                    break
            
            if skill_matrix_start_idx is not None:
                # Find the end of the skill matrix section (before EDUCATION or WORK EXPERIENCE)
                end_idx = skill_matrix_start_idx + 1
                while end_idx < len(doc.paragraphs):
                    text = doc.paragraphs[end_idx].text.strip()
                    if not text:
                        end_idx += 1
                        continue
                    # Stop at next major section
                    if any(h in text for h in ("EDUCATION", "WORK EXPERIENCE")):
                        break
                    end_idx += 1
                
                # Remove all old skill paragraphs
                for p in doc.paragraphs[skill_matrix_start_idx + 1:end_idx]:
                    p._element.getparent().remove(p._element)
                
                # Check if first category header exists in template - if so, replace it; if not, insert it
                first_header_exists = False
                if skill_matrix_start_idx < len(doc.paragraphs) and category_headers:
                    existing_text = doc.paragraphs[skill_matrix_start_idx].text.strip()
                    if category_headers[0] in existing_text:
                        first_header_exists = True
                        # Replace the existing header text
                        doc.paragraphs[skill_matrix_start_idx].clear()
                        doc.paragraphs[skill_matrix_start_idx].add_run(category_headers[0])
                        for run in doc.paragraphs[skill_matrix_start_idx].runs:
                            run.font.size = Pt(10)
                
                # Parse the LLM response and insert structured content
                lines = skill_matrix_content.split('\n')
                current_category = None
                insert_pos = skill_matrix_start_idx + 1
                inserted_categories = set()  # Track which categories have been inserted
                
                # Add first category to inserted set if it already exists, or insert it if it doesn't
                if category_headers:
                    if first_header_exists:
                        inserted_categories.add(category_headers[0])
                        current_category = category_headers[0]
                    else:
                        # Insert the first category header
                        if insert_pos < len(doc.paragraphs):
                            anchor = doc.paragraphs[insert_pos]
                            new_p = anchor.insert_paragraph_before(category_headers[0])
                        else:
                            new_p = doc.add_paragraph(category_headers[0])
                        for run in new_p.runs:
                            run.font.size = Pt(10)
                        inserted_categories.add(category_headers[0])
                        current_category = category_headers[0]
                        insert_pos += 1
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if this line matches any of the generated category headers
                    is_category_header = False
                    category_name = None
                    
                    # Check against all generated category headers
                    for header in category_headers:
                        # Exact match or starts with header and has colon
                        if line == header or line.startswith(header + ":"):
                            category_name = header
                            is_category_header = True
                            break
                        # Also check if line contains the header (for flexibility)
                        elif header in line and (":" in line or len(line) < len(header) + 20):
                            category_name = header
                            is_category_header = True
                            break
                    
                    if is_category_header and category_name:
                        current_category = category_name
                        # Only insert if not already inserted
                        if category_name not in inserted_categories:
                            inserted_categories.add(category_name)
                            # Insert category header
                            if insert_pos < len(doc.paragraphs):
                                anchor = doc.paragraphs[insert_pos]
                                new_p = anchor.insert_paragraph_before(category_name)
                            else:
                                new_p = doc.add_paragraph(category_name)
                            for run in new_p.runs:
                                run.font.size = Pt(10)
                            insert_pos += 1
                        continue
                    
                    # This is a bullet point (starts with - or • or number, or is a sentence)
                    # Only process if we have a current category
                    if current_category:
                        # Check if it's a bullet (starts with marker) or a regular sentence
                        is_bullet = line.startswith('-') or line.startswith('•') or line.startswith('*') or \
                                   (line[0].isdigit() and '.' in line[:3])
                        
                        if is_bullet:
                            # Remove bullet markers
                            bullet_text = line.lstrip('- •*0123456789. ').strip()
                        else:
                            # Might be a sentence without bullet marker - treat as content
                            bullet_text = line.strip()
                        
                        if bullet_text and len(bullet_text) > 10:  # Only add substantial content
                            # Insert as bullet paragraph
                            if insert_pos < len(doc.paragraphs):
                                anchor = doc.paragraphs[insert_pos]
                                new_p = anchor.insert_paragraph_before(bullet_text)
                            else:
                                new_p = doc.add_paragraph(bullet_text)
                            # Try to apply bullet style
                            try:
                                # Look for bullet style
                                bullet_style = None
                                for style in doc.styles:
                                    if "bullet" in style.name.lower() or "list" in style.name.lower():
                                        bullet_style = style.name
                                        break
                                if bullet_style:
                                    new_p.style = bullet_style
                            except:
                                pass
                            for run in new_p.runs:
                                run.font.size = Pt(10)
                            insert_pos += 1
        except Exception as e:
            # If generation fails, fall back to simple skill list
            pass

    # === 5. Education ===
    edu_start = None
    for i, p in enumerate(doc.paragraphs):
        if "EDUCATION" in p.text:
            edu_start = i + 1
            break
    if edu_start:
        # collect all existing education paragraphs until a blank or next heading
        end_idx = edu_start
        while end_idx < len(doc.paragraphs):
            text = doc.paragraphs[end_idx].text.strip()
            if not text or "WORK EXPERIENCE" in text or doc.paragraphs[end_idx].style.name.startswith("Heading"):
                break
            end_idx += 1

        # Anchor is the paragraph where new entries will be inserted before
        anchor = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

        # Remove old education paragraphs (between edu_start and end_idx)
        for p in doc.paragraphs[edu_start:end_idx]:
            p._element.getparent().remove(p._element)

        # Insert new education lines in reverse order to preserve order
        if anchor:
            for edu in reversed(data.get("education", [])):
                line = f"{edu['degree']}, {edu['school']}, {edu['year']}"
                new_p = anchor.insert_paragraph_before(line)
                for run in new_p.runs:
                    run.font.size = Pt(10)

    # === 6. Work Experience – delete old, add new with exact same style ===
    exp_start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "WORK EXPERIENCE" in p.text:
            exp_start_idx = i + 1
            break

    if exp_start_idx:
        # remove old jobs (delete paragraphs instead of clearing to avoid empty bullet lines)
        while exp_start_idx < len(doc.paragraphs):
            p = doc.paragraphs[exp_start_idx]
            if p.style.name.startswith("Heading"):
                break
            p._element.getparent().remove(p._element)

        # Try to detect a bullet/list style that exists in this template
        bullet_style_name = None
        try:
            for s in doc.styles:
                if "bullet" in str(getattr(s, "name", "")).lower():
                    bullet_style_name = s.name
                    break
        except Exception:
            bullet_style_name = None

        # Append new experience entries at the end of the document to avoid
        # low-level XML manipulation that can return None on some templates.
        for job in data["experience"]:
            # Job title + dates
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{job['title']} [{job['dates']}]")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Company
            comp_para = doc.add_paragraph()
            comp_run = comp_para.add_run(job['company'])
            comp_run.font.size = Pt(11)
            comp_run.font.color.rgb = RGBColor(0, 112, 192)

            # Bullets – use an existing bullet style if available, otherwise plain paragraphs
            for bullet in job['bullets']:
                # Skip empty / whitespace-only bullets to avoid stray "•" lines
                if not bullet or not str(bullet).strip():
                    continue
                text = str(bullet).strip()
                if bullet_style_name:
                    bullet_para = doc.add_paragraph(text, style=bullet_style_name)
                else:
                    bullet_para = doc.add_paragraph(f"• {text}")
                for run in bullet_para.runs:
                    run.font.size = Pt(10)

    # Save
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

def generate_resume(candidate_resume_file):
    raw_text = read_any_resume(candidate_resume_file)
    data = extract_with_llama70b(raw_text)

    # If summary is empty, generate one based on the resume content
    if not data.get("summary") or not str(data.get("summary", "")).strip():
        generated_summary = generate_summary_from_resume(
            raw_text,
            data.get("experience", []),
            data.get("education", []),
            data.get("skills", "")
        )
        if generated_summary:
            data["summary"] = generated_summary

    # Load your original template (bundled with the script)
    with open("main_resume.docx", "rb") as f:
        template_bytes = f.read()

    new_docx_bytes = apply_ATS_template(template_bytes, data)

    # Gradio's File output expects a path-like, not raw bytes.
    # Write the generated DOCX to a temporary file and return its path.
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(new_docx_bytes)
        tmp_path = tmp.name

    return tmp_path, data

# ========================== GRADIO UI ==========================
with gr.Blocks(title="ATS-Style Resume Cloner") as demo:
    gr.Markdown("# ATS Resume Cloner\n"
                "Drop any resume (PDF/DOCX/TXT) → get a perfect copy in **your exact beautiful style** instantly")

    candidate = gr.File(label="Candidate's Resume (any format)", file_types=[".pdf",".docx",".txt"])
    btn = gr.Button("Generate My ATS-Style Resume", variant="primary", size="lg")

    out_docx = gr.File(label="Your new perfect resume.docx")
    out_json = gr.JSON(label="Extracted data (for checking)")

    btn.click(generate_resume, inputs=candidate, outputs=[out_docx, out_json])

demo.launch(share=False)